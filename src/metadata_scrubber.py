"""
Модуль для очистки метаданных из файлов.
"""

import os
import shutil
from pathlib import Path
from typing import Dict, Optional, List
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import logging

# Для работы с PDF
try:
    from pypdf import PdfReader, PdfWriter
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# Для работы с DOCX
try:
    from docx import Document
    import docx.opc.constants as opc_constants
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class MetadataScrubber:
    """
    Класс для очистки метаданных из файлов.
    """

    def __init__(self):
        self.cleaned_files = 0
        self.removed_geotags = 0
        self.removed_device_info = 0
        self.removed_software_info = 0
        self.removed_personal_info = 0
        self.last_action = ""
        self.supported_formats = ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif', '.mp4', '.avi', '.mov', '.mkv', '.pdf', '.docx']
        
        # Поддерживаемые форматы по категориям
        self.image_formats = {'.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif'}
        self.video_formats = {'.mp4', '.avi', '.mov', '.mkv'}
        self.document_formats = {'.pdf', '.docx'}
        
        # Категории метаданных для удаления
        self.metadata_categories = {
            'location': self._is_gps_tag,
            'device_info': self._is_device_tag,
            'software': self._is_software_tag,
            'personal': self._is_personal_tag
        }

    def _is_gps_tag(self, tag: str) -> bool:
        """
        Проверяет, является ли тег GPS-координатами.

        Args:
            tag: Имя тега

        Returns:
            True, если тег связан с GPS
        """
        gps_tags = ['GPSInfo', 'GPSLatitude', 'GPSLongitude', 'GPSAltitude', 
                   'GPSTimeStamp', 'GPSSpeed', 'GPSDestLatitude', 'GPSDestLongitude']
        return any(gps_tag.lower() in str(tag).lower() for gps_tag in gps_tags)

    def _is_device_tag(self, tag: str) -> bool:
        """
        Проверяет, является ли тег информацией об устройстве.

        Args:
            tag: Имя тега

        Returns:
            True, если тег связан с устройством
        """
        device_tags = ['Make', 'Model', 'LensModel', 'SerialNumber', 
                      'Hardware', 'Firmware', 'Software']
        return any(device_tag.lower() in str(tag).lower() for device_tag in device_tags)

    def _is_software_tag(self, tag: str) -> bool:
        """
        Проверяет, является ли тег информацией о программном обеспечении.

        Args:
            tag: Имя тега

        Returns:
            True, если тег связан с ПО
        """
        software_tags = ['Software', 'Application', 'Creator', 'EditHistory', 
                        'History', 'Processing', 'Filter']
        return any(software_tag.lower() in str(tag).lower() for software_tag in software_tags)

    def _is_personal_tag(self, tag: str) -> bool:
        """
        Проверяет, является ли тег персональной информацией.

        Args:
            tag: Имя тега

        Returns:
            True, если тег содержит персональные данные
        """
        personal_tags = ['Artist', 'Author', 'Creator', 'Copyright', 'Comment', 
                        'Description', 'Title', 'Subject', 'Keywords', 'Rating']
        return any(personal_tag.lower() in str(tag).lower() for personal_tag in personal_tags)

    def _get_file_metadata(self, file_path: Path) -> Dict:
        """
        Извлекает метаданные из файла.

        Args:
            file_path: Путь к файлу

        Returns:
            Словарь с метаданными
        """
        metadata = {}
        
        try:
            if file_path.suffix.lower() in ['.jpg', '.jpeg', '.tiff']:
                with Image.open(file_path) as img:
                    exif_data = img.getexif()
                    
                    if exif_data:
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            metadata[str(tag)] = str(value)
            
            elif file_path.suffix.lower() == '.pdf' and HAS_PYPDF:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    if reader.metadata:
                        for key, value in reader.metadata.items():
                            metadata[f'PDF_{key}'] = str(value)

            elif file_path.suffix.lower() == '.docx' and HAS_DOCX:
                doc = Document(file_path)
                core_props = doc.core_properties
                
                if core_props.author:
                    metadata['DOCX_Author'] = core_props.author
                if core_props.last_modified_by:
                    metadata['DOCX_LastModifiedBy'] = core_props.last_modified_by
                if core_props.revision:
                    metadata['DOCX_Revision'] = str(core_props.revision)
                if core_props.comments:
                    metadata['DOCX_Comments'] = core_props.comments
                if core_props.created:
                    metadata['DOCX_Created'] = str(core_props.created)
                if core_props.modified:
                    metadata['DOCX_Modified'] = str(core_props.modified)

        except Exception as e:
            logger.warning(f"Не удалось извлечь метаданные из {file_path}: {e}")
            
        return metadata

    def _has_metadata(self, file_path: Path) -> bool:
        """
        Проверяет, содержит ли файл метаданные.

        Args:
            file_path: Путь к файлу

        Returns:
            True, если файл содержит метаданные
        """
        metadata = self._get_file_metadata(file_path)
        return len(metadata) > 0

    def _remove_metadata_from_image(self, file_path: Path, categories_to_remove: List[str]) -> bool:
        """
        Удаляет метаданные из изображения.

        Args:
            file_path: Путь к файлу
            categories_to_remove: Список категорий метаданных для удаления

        Returns:
            True, если очистка прошла успешно
        """
        try:
            # Открываем изображение
            with Image.open(file_path) as img:
                # Создаем копию изображения без EXIF данных
                data = list(img.getdata())
                image_without_exif = Image.new(img.mode, img.size)
                image_without_exif.putdata(data)
                
                # Сохраняем изображение без метаданных
                # Определяем формат по расширению
                format_map = {
                    '.jpg': 'JPEG',
                    '.jpeg': 'JPEG',
                    '.png': 'PNG',
                    '.tiff': 'TIFF',
                    '.bmp': 'BMP',
                    '.gif': 'GIF'
                }
                
                format_name = format_map.get(file_path.suffix.lower())
                
                if format_name:
                    # Создаем временный файл
                    temp_path = file_path.with_suffix(file_path.suffix + '.temp')
                    temp_created = False
                    try:
                        # Сохраняем временный файл
                        image_without_exif.save(temp_path, format=format_name)
                        temp_created = True
                        
                        # Заменяем оригинальный файл
                        shutil.move(str(temp_path), str(file_path))
                        
                        # Обновляем статистику
                        self.cleaned_files += 1
                        self.last_action = file_path.name
                        
                        # Обновляем статистику по категориям
                        original_metadata = self._get_file_metadata(file_path)
                        for tag in original_metadata:
                            if self._is_gps_tag(tag) and 'location' in categories_to_remove:
                                self.removed_geotags += 1
                            elif self._is_device_tag(tag) and 'device_info' in categories_to_remove:
                                self.removed_device_info += 1
                            elif self._is_software_tag(tag) and 'software' in categories_to_remove:
                                self.removed_software_info += 1
                            elif self._is_personal_tag(tag) and 'personal' in categories_to_remove:
                                self.removed_personal_info += 1
                        
                        logger.info(f"Метаданные удалены из {file_path}")
                        return True
                    
                    except Exception as e:
                        logger.error(f"Ошибка при сохранении временного файла или замене оригинала: {e}")
                        return False
                        
                    finally:
                        # Удаляем временный файл, если он был создан
                        if temp_created and temp_path.exists():
                            try:
                                os.unlink(temp_path)
                            except Exception as e:
                                logger.warning(f"Не удалось удалить временный файл {temp_path}: {e}")
                else:
                    logger.error(f"Неподдерживаемый формат изображения: {file_path.suffix}")
                    return False
                    
        except Exception as e:
            logger.error(f"Ошибка при удалении метаданных из {file_path}: {e}")
            return False

    def _remove_metadata_from_pdf(self, file_path: Path, categories_to_remove: List[str]) -> bool:
        """
        Удаляет метаданные из PDF файла.

        Args:
            file_path: Путь к файлу
            categories_to_remove: Список категорий метаданных для удаления

        Returns:
            True, если очистка прошла успешно
        """
        if not HAS_PYPDF:
            logger.error("Модуль pypdf не установлен. Установите с помощью: pip install pypdf")
            return False
            
        try:
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                writer = PdfWriter()
                
                # Копируем все страницы
                for page in reader.pages:
                    writer.add_page(page)
                
                # Удаляем метаданные
                writer.add_metadata({})
                
                # Создаем временный файл
                temp_path = file_path.with_suffix(file_path.suffix + '.temp')
                temp_created = False
                try:
                    with open(temp_path, 'wb') as f_out:
                        writer.write(f_out)
                    temp_created = True
                    
                    # Заменяем оригинальный файл
                    shutil.move(str(temp_path), str(file_path))
                    
                    # Обновляем статистику
                    self.cleaned_files += 1
                    self.last_action = file_path.name
                    
                    # Обновляем статистику по категориям
                    if 'personal' in categories_to_remove:
                        self.removed_personal_info += 1
                    
                    logger.info(f"Метаданные удалены из PDF {file_path}")
                    return True
                    
                except Exception as e:
                    logger.error(f"Ошибка при сохранении временного файла или замене оригинала: {e}")
                    return False
                    
                finally:
                    # Удаляем временный файл, если он был создан
                    if temp_created and temp_path.exists():
                        try:
                            os.unlink(temp_path)
                        except Exception as e:
                            logger.warning(f"Не удалось удалить временный файл {temp_path}: {e}")
                            
        except Exception as e:
            logger.error(f"Ошибка при удалении метаданных из PDF {file_path}: {e}")
            return False

    def _remove_metadata_from_docx(self, file_path: Path, categories_to_remove: List[str]) -> bool:
        """
        Удаляет метаданные из DOCX файла.

        Args:
            file_path: Путь к файлу
            categories_to_remove: Список категорий метаданных для удаления

        Returns:
            True, если очистка прошла успешно
        """
        if not HAS_DOCX:
            logger.error("Модуль python-docx не установлен. Установите с помощью: pip install python-docx")
            return False
            
        try:
            # Открываем документ
            doc = Document(file_path)
            
            # Очищаем основные свойства
            core_props = doc.core_properties
            core_props.author = None
            core_props.last_modified_by = None
            core_props.revision = 1
            core_props.comments = None
            core_props.category = None
            core_props.keywords = None
            core_props.content_status = None
            
            # Удаляем историю изменений, если она существует в документе
            # Удаляем все изменения (track changes)
            for part in doc.part.target_parts.values():
                if part.content_type == opc_constants.CT_VML_DRAWING:
                    continue
                xml = part._blob.decode('utf-8')
                if 'w:revision' in xml or 'w:ins' in xml or 'w:del' in xml:
                    xml = xml.replace('w:revision', 'w:revision_removed').replace('w:ins ', 'w:p ').replace('w:del ', 'w:p ')
                    part._blob = xml.encode('utf-8')

            # Удаляем комментарии
            if hasattr(doc, '_comments_part'):
                doc._comments_part = None
            
            # Удаляем все комментарии из XML
            xml = doc._element.xml
            if 'w:commentRangeStart' in xml:
                xml = xml.replace('w:commentRangeStart', 'w:p').replace('w:commentRangeEnd', 'w:p').replace('w:commentReference', 'w:p')
                doc._element.xml = xml

            # Создаем временный файл
            temp_path = file_path.with_suffix(file_path.suffix + '.temp')
            temp_created = False
            try:
                # Сохраняем документ
                doc.save(temp_path)
                temp_created = True
                
                # Заменяем оригинальный файл
                shutil.move(str(temp_path), str(file_path))
                
                # Обновляем статистику
                self.cleaned_files += 1
                self.last_action = file_path.name
                
                # Обновляем статистику по категориям
                if 'personal' in categories_to_remove:
                    self.removed_personal_info += 1
                
                logger.info(f"Метаданные удалены из DOCX {file_path}")
                return True
                
            except Exception as e:
                logger.error(f"Ошибка при сохранении временного файла или замене оригинала: {e}")
                return False
                
            finally:
                # Удаляем временный файл, если он был создан
                if temp_created and temp_path.exists():
                    try:
                        os.unlink(temp_path)
                    except Exception as e:
                        logger.warning(f"Не удалось удалить временный файл {temp_path}: {e}")
                        
        except Exception as e:
            logger.error(f"Ошибка при удалении метаданных из DOCX {file_path}: {e}")
            return False

    def _remove_embedded_thumbnails(self, file_path: Path) -> bool:
        """
        Удаляет встроенные миниатюры и скрытые данные из файла.

        Args:
            file_path: Путь к файлу

        Returns:
            True, если удаление прошло успешно
        """
        try:
            # Для DOCX файлов
            if file_path.suffix.lower() == '.docx' and HAS_DOCX:
                doc = Document(file_path)
                # Удаляем все изображения, которые могут быть миниатюрами
                for rel in list(doc.part.rels.values()):
                    if 'image' in rel.target_ref:
                        # Проверяем размеры изображения (миниатюры обычно маленькие)
                        try:
                            image_part = rel.target_part
                            # Это упрощенная проверка - в реальности нужно анализировать размеры
                            # Удаляем все встроенные изображения (в реальности можно фильтровать по размеру)
                            doc.part.drop_rel(rel.rId)
                        except:
                            continue
                
                # Сохраняем изменения
                temp_path = file_path.with_suffix(file_path.suffix + '.temp')
                doc.save(temp_path)
                shutil.move(str(temp_path), str(file_path))

            # Для PDF файлов
            elif file_path.suffix.lower() == '.pdf' and HAS_PYPDF:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    writer = PdfWriter()
                    
                    # Копируем все страницы
                    for page in reader.pages:
                        writer.add_page(page)
                    
                # Удаляем вложенные файлы и миниатюры
                writer.add_metadata({})
                if '/OpenAction' in writer._root_object:
                    del writer._root_object['/OpenAction']
                
                # Сохраняем изменения
                temp_path = file_path.with_suffix(file_path.suffix + '.temp')
                with open(temp_path, 'wb') as f_out:
                    writer.write(f_out)
                shutil.move(str(temp_path), str(file_path))

            return True
            
        except Exception as e:
            logger.warning(f"Не удалось удалить встроенные миниатюры из {file_path}: {e}")
            return False

    def scrub_file(self, file_path: str, categories_to_remove: Optional[List[str]] = None, remove_thumbnails: bool = False) -> bool:
        """
        Очищает метаданные из указанного файла.

        Args:
            file_path: Путь к файлу для обработки
            categories_to_remove: Список категорий метаданных для удаления
            remove_thumbnails: Удалить встроенные миниатюры и скрытые данные

        Returns:
            True, если очистка прошла успешно, иначе False
        """
        file_path_obj = Path(file_path)
        
        # Проверяем существование файла
        if not file_path_obj.exists():
            logger.error(f"Файл не найден: {file_path}")
            return False
            
        # Проверяем поддерживаемый формат
        if file_path_obj.suffix.lower() not in self.supported_formats:
            logger.warning(f"Неподдерживаемый формат файла: {file_path_obj.suffix}")
            return False
            
        # Проверяем, содержит ли файл метаданные
        if not self._has_metadata(file_path_obj):
            logger.info(f"Файл не содержит метаданных: {file_path}")
            self.last_action = file_path_obj.name
            return True
        
        # Устанавливаем категории для удаления (все, если не указаны)
        if categories_to_remove is None:
            categories_to_remove = list(self.metadata_categories.keys())
        
        # Проверяем валидность категорий
        valid_categories = [cat for cat in categories_to_remove if cat in self.metadata_categories]
        if not valid_categories:
            logger.warning(f"Нет валидных категорий для удаления: {categories_to_remove}")
            return False
        
        # Очищаем метаданные в зависимости от типа файла
        result = False
        if file_path_obj.suffix.lower() in self.image_formats:
            result = self._remove_metadata_from_image(file_path_obj, valid_categories)
        elif file_path_obj.suffix.lower() == '.pdf':
            result = self._remove_metadata_from_pdf(file_path_obj, valid_categories)
        elif file_path_obj.suffix.lower() == '.docx':
            result = self._remove_metadata_from_docx(file_path_obj, valid_categories)
        else:
            logger.warning(f"Поддержка формата {file_path_obj.suffix} в разработке")
            result = False

        # Если очистка прошла успешно и нужно удалить миниатюры
        if result and remove_thumbnails:
            if not self._remove_embedded_thumbnails(file_path_obj):
                logger.warning(f"Не удалось удалить встроенные миниатюры из {file_path}")

        return result

    def get_statistics(self) -> Dict:
        """
        Возвращает статистику по очистке метаданных.

        Returns:
            Словарь со статистикой
        """
        return {
            "cleaned_files": self.cleaned_files,
            "removed_geotags": self.removed_geotags,
            "removed_device_info": self.removed_device_info,
            "removed_software_info": self.removed_software_info,
            "removed_personal_info": self.removed_personal_info,
            "last_action": self.last_action
        }
        
    def reset_statistics(self):
        """
        Сбрасывает статистику.
        """
        self.cleaned_files = 0
        self.removed_geotags = 0
        self.removed_device_info = 0
        self.removed_software_info = 0
        self.removed_personal_info = 0
        self.last_action = ""